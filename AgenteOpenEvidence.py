import os
import faiss
import numpy as np
from typing import List, Dict, Any, Optional
from openai import AzureOpenAI
import trafilatura
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from dotenv import load_dotenv
import PyPDF2
import docx
import pandas as pd
from googleapiclient.discovery import build
import pickle
import time
from urllib.parse import urlparse
import logging

load_dotenv()

# Configurar logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class MedicalRAGSystem:
    # Dominios médicos confiables
    TRUSTED_MEDICAL_DOMAINS = {
        'pubmed.ncbi.nlm.nih.gov',
        'www.ncbi.nlm.nih.gov',
        'www.nejm.org',
        'jamanetwork.com',
        'www.thelancet.com',
        'www.bmj.com',
        'www.nature.com',
        'science.sciencemag.org',
        'www.cell.com',
        'www.cochranelibrary.com',
        'www.who.int',
        'www.cdc.gov',
        'www.fda.gov',
        'www.ema.europa.eu',
        'www.mayoclinic.org',
        'www.clevelandclinic.org',
        'www.uptodate.com',
        'www.medlineplus.gov',
        'emedicine.medscape.com',
        'www.webmd.com',
        'www.healthline.com'
    }
    
    MEDICAL_KEYWORDS = {
        'disease', 'treatment', 'diagnosis', 'symptom', 'therapy', 'medication',
        'patient', 'clinical', 'medical', 'health', 'drug', 'syndrome',
        'pathology', 'pharmacology', 'epidemiology', 'surgery', 'hospital',
        'enfermedad', 'tratamiento', 'diagnóstico', 'síntoma', 'terapia',
        'medicamento', 'paciente', 'clínico', 'médico', 'salud', 'fármaco'
    }

    def __init__(self, 
                openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
                openai_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
                deployment_name=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
                api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
                embedding_deployment=os.getenv("AZURE_EMBEDDING_DEPLOYMENT"),
                google_api_key=os.getenv("GOOGLE_SEARCH_API_KEY"),
                google_cse_id=os.getenv("GOOGLE_CSE_ID"),
                storage_path: str = "medical_knowledge_base",
                chunk_size: int = 800,
                chunk_overlap: int = 150,
                embedding_dimension: int = 3072):  # text-embedding-3-large dimension
        
        self.openai_client = AzureOpenAI(
            api_key=openai_api_key,
            api_version=api_version,
            azure_endpoint=openai_endpoint
        )
        self.deployment_name = deployment_name
        self.embedding_deployment = embedding_deployment or "text-embedding-3-large"
        self.storage_path = storage_path
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_dimension = embedding_dimension
        
        # Configuración para Google Search
        self.google_api_key = google_api_key
        self.google_cse_id = google_cse_id
        
        # Inicializar FAISS
        self.index = None
        self.documents_metadata = []
        self._load_or_create_index()
        
        # Configurar Selenium para scraping avanzado
        self.chrome_options = Options()
        self.chrome_options.add_argument('--headless')
        self.chrome_options.add_argument('--no-sandbox')
        self.chrome_options.add_argument('--disable-dev-shm-usage')
        self.chrome_options.add_argument('--disable-gpu')
        
        # Prompt optimizado para medicina
        self.system_prompt = """Eres un asistente médico especializado que analiza literatura científica y documentos médicos para proporcionar información precisa y basada en evidencia.

        INSTRUCCIONES CRÍTICAS:
        1. SOLO utiliza información de fuentes médicas confiables
        2. Proporciona respuestas detalladas y técnicamente precisas cuando tengas evidencia suficiente
        3. Siempre cita las fuentes específicas con nivel de evidencia cuando sea posible
        4. Si la información es contradictoria entre estudios, analízala críticamente
        5. Menciona las limitaciones de los estudios o la evidencia disponible
        6. NUNCA des consejos médicos directos - siempre recomienda consultar con un profesional de la salud
        7. Clasifica el nivel de evidencia (Nivel I-V, Grado A-D según sea apropiado)
        8. Estructura las respuestas de manera clara para profesionales de la salud

        IMPORTANTE: Esta información es solo para fines educativos y de investigación. Siempre consulte con un profesional médico calificado."""

        logger.info(f"🏥 Sistema RAG Médico inicializado con Azure Embedding: {self.embedding_deployment}")

    def _load_or_create_index(self):
        """Cargar índice FAISS existente o crear uno nuevo"""
        try:
            index_path = os.path.join(self.storage_path, 'faiss_index.bin')
            metadata_path = os.path.join(self.storage_path, 'metadata.pkl')
            
            if os.path.exists(index_path) and os.path.exists(metadata_path):
                # Cargar índice existente
                self.index = faiss.read_index(index_path)
                with open(metadata_path, 'rb') as f:
                    self.documents_metadata = pickle.load(f)
                logger.info(f"📚 Índice FAISS cargado: {len(self.documents_metadata)} documentos")
            else:
                # Crear nuevo índice
                os.makedirs(self.storage_path, exist_ok=True)
                self.index = faiss.IndexFlatIP(self.embedding_dimension)  # Inner Product para cosine similarity
                self.documents_metadata = []
                logger.info("🆕 Nuevo índice FAISS creado")
                
        except Exception as e:
            logger.error(f"❌ Error inicializando FAISS: {e}")
            # Fallback: crear nuevo índice
            os.makedirs(self.storage_path, exist_ok=True)
            self.index = faiss.IndexFlatIP(self.embedding_dimension)
            self.documents_metadata = []

    def _save_index(self):
        """Guardar índice FAISS y metadata"""
        try:
            index_path = os.path.join(self.storage_path, 'faiss_index.bin')
            metadata_path = os.path.join(self.storage_path, 'metadata.pkl')
            
            faiss.write_index(self.index, index_path)
            with open(metadata_path, 'wb') as f:
                pickle.dump(self.documents_metadata, f)
            logger.info("💾 Índice FAISS guardado exitosamente")
            
        except Exception as e:
            logger.error(f"❌ Error guardando índice: {e}")

    def get_batch_embeddings(self, texts: List[str], batch_size: int = 100) -> List[List[float]]:
        """Obtener embeddings en lotes para optimizar rendimiento"""
        embeddings = []
        
        try:
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                logger.info(f"🔄 Procesando batch {i//batch_size + 1}/{(len(texts) + batch_size - 1)//batch_size}")
                
                response = self.openai_client.embeddings.create(
                    input=batch,
                    model=self.embedding_deployment
                )
                
                batch_embeddings = [item.embedding for item in response.data]
                embeddings.extend(batch_embeddings)
                
                # Pequeña pausa para evitar rate limiting
                if i + batch_size < len(texts):
                    time.sleep(0.1)
                    
            return embeddings
            
        except Exception as e:
            logger.error(f"❌ Error obteniendo embeddings por lotes: {e}")
            return []

    def _is_trusted_medical_source(self, url: str) -> bool:
        """Verificar si la URL es de una fuente médica confiable"""
        try:
            domain = urlparse(url).netloc.lower()
            return any(trusted in domain for trusted in self.TRUSTED_MEDICAL_DOMAINS)
        except:
            return False

    def _has_medical_content(self, text: str) -> bool:
        """Verificar si el texto contiene contenido médico relevante"""
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in self.MEDICAL_KEYWORDS if keyword in text_lower)
        return keyword_count >= 3  # Al menos 3 términos médicos

    def extract_content_trafilatura(self, url: str) -> Optional[str]:
        """Extraer contenido usando trafilatura (más eficiente)"""
        try:
            # Verificar si es fuente confiable
            if not self._is_trusted_medical_source(url):
                logger.warning(f"⚠️ Fuente no confiable ignorada: {url}")
                return None
                
            downloaded = trafilatura.fetch_url(url)
            if downloaded:
                content = trafilatura.extract(downloaded, 
                                            include_comments=False,
                                            include_tables=True,
                                            include_formatting=True)
                
                if content and self._has_medical_content(content):
                    logger.info(f"✅ Contenido médico extraído con trafilatura: {url}")
                    return content
                else:
                    logger.warning(f"⚠️ Contenido no relevante médicamente: {url}")
                    return None
            return None
            
        except Exception as e:
            logger.error(f"❌ Error con trafilatura en {url}: {e}")
            return None

    def extract_content_selenium(self, url: str) -> Optional[str]:
        """Extraer contenido usando Selenium (para sitios con JS)"""
        driver = None
        try:
            # Verificar fuente confiable
            if not self._is_trusted_medical_source(url):
                logger.warning(f"⚠️ Fuente no confiable ignorada: {url}")
                return None
                
            driver = webdriver.Chrome(options=self.chrome_options)
            driver.get(url)
            
            # Esperar a que cargue el contenido
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )
            
            # Extraer texto del cuerpo principal
            content_selectors = [
                'main', 'article', '.content', '#content', 
                '.main-content', '.article-content', '.abstract',
                '.full-text', '.article-body'
            ]
            
            content = ""
            for selector in content_selectors:
                try:
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        content = elements[0].get_attribute('innerText')
                        break
                except:
                    continue
            
            # Si no encuentra selectores específicos, usar body
            if not content:
                content = driver.find_element(By.TAG_NAME, "body").get_attribute('innerText')
            
            if content and self._has_medical_content(content):
                logger.info(f"✅ Contenido médico extraído con Selenium: {url}")
                return content[:8000]  # Limitar tamaño
            else:
                logger.warning(f"⚠️ Contenido no relevante médicamente: {url}")
                return None
                
        except Exception as e:
            logger.error(f"❌ Error con Selenium en {url}: {e}")
            return None
        finally:
            if driver:
                driver.quit()

    def medical_search(self, query: str, num_results: int = 5) -> List[Dict[str, Any]]:
        """Búsqueda especializada en fuentes médicas"""
        if not self.google_api_key or not self.google_cse_id:
            logger.warning("⚠️ Google Search API no configurada")
            return []
        
        try:
            service = build("customsearch", "v1", developerKey=self.google_api_key)
            
            # Búsqueda específica en sitios médicos
            medical_query = f"{query} site:pubmed.ncbi.nlm.nih.gov OR site:nejm.org OR site:jamanetwork.com OR site:thelancet.com OR site:bmj.com OR site:who.int OR site:cdc.gov"
            
            result = service.cse().list(
                q=medical_query,
                cx=self.google_cse_id,
                num=num_results
            ).execute()
            
            search_results = []
            
            if 'items' in result:
                for item in result['items']:
                    url = item['link']
                    
                    # Doble verificación de fuente confiable
                    if not self._is_trusted_medical_source(url):
                        continue
                    
                    # Intentar primero con trafilatura (más eficiente)
                    content = self.extract_content_trafilatura(url)
                    
                    # Si falla, intentar con Selenium
                    if not content:
                        content = self.extract_content_selenium(url)
                    
                    if content:
                        search_results.append({
                            'title': item['title'],
                            'link': url,
                            'snippet': item['snippet'],
                            'content': content,
                            'source_type': 'MEDICAL_WEB'
                        })
            
            logger.info(f"🏥 Encontrados {len(search_results)} resultados médicos para: {query}")
            return search_results
            
        except Exception as e:
            logger.error(f"❌ Error en búsqueda médica: {e}")
            return []

    def process_file(self, file_path: str, filename: str) -> str:
        """Procesa diferentes tipos de archivos médicos"""
        try:
            file_extension = filename.lower().split('.')[-1]
            content = ""
            
            if file_extension == 'pdf':
                content = self._process_pdf(file_path)
            elif file_extension in ['docx', 'doc']:
                content = self._process_docx(file_path)
            elif file_extension == 'txt':
                content = self._process_txt(file_path)
            elif file_extension in ['csv']:
                content = self._process_csv(file_path)
            elif file_extension in ['xlsx', 'xls']:
                content = self._process_excel(file_path)
            else:
                raise ValueError(f"Tipo de archivo no soportado: {file_extension}")
            
            # Verificar relevancia médica
            if not self._has_medical_content(content):
                logger.warning(f"⚠️ Archivo sin contenido médico relevante: {filename}")
                return ""
            
            return content
            
        except Exception as e:
            logger.error(f"❌ Error procesando archivo {filename}: {e}")
            return ""

    def _process_pdf(self, file_path: str) -> str:
        """Procesa archivos PDF optimizado para documentos médicos"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n--- Página {page_num + 1} ---\n{page_text}\n"
            return content
        except Exception as e:
            logger.error(f"❌ Error procesando PDF: {e}")
            return ""

    def _process_docx(self, file_path: str) -> str:
        """Procesa archivos DOCX"""
        try:
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Procesar tablas también
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    content += " | ".join(row_data) + "\n"
                    
            return content
        except Exception as e:
            logger.error(f"❌ Error procesando DOCX: {e}")
            return ""

    def _process_txt(self, file_path: str) -> str:
        """Procesa archivos TXT"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            return ""
        except Exception as e:
            logger.error(f"❌ Error procesando TXT: {e}")
            return ""

    def _process_csv(self, file_path: str) -> str:
        """Procesa archivos CSV de datos médicos"""
        try:
            df = pd.read_csv(file_path)
            
            content = f"DATASET MÉDICO - Información estructurada:\n\n"
            content += f"Dimensiones: {df.shape[0]} registros x {df.shape[1]} variables\n\n"
            content += f"Variables: {', '.join(df.columns.tolist())}\n\n"
            
            # Estadísticas para variables numéricas
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content += "ESTADÍSTICAS DESCRIPTIVAS:\n"
                content += df[numeric_cols].describe().to_string() + "\n\n"
            
            # Muestra de datos
            content += "MUESTRA DE DATOS (primeras 10 filas):\n"
            content += df.head(10).to_string() + "\n\n"
            
            # Información sobre variables categóricas
            categorical_cols = df.select_dtypes(include=['object']).columns
            if len(categorical_cols) > 0:
                content += "VARIABLES CATEGÓRICAS:\n"
                for col in categorical_cols:
                    unique_vals = df[col].value_counts().head(10)
                    content += f"{col}:\n{unique_vals.to_string()}\n\n"
            
            return content
            
        except Exception as e:
            logger.error(f"❌ Error procesando CSV: {e}")
            return ""

    def _process_excel(self, file_path: str) -> str:
        """Procesa archivos Excel de datos médicos"""
        try:
            excel_file = pd.ExcelFile(file_path)
            content = f"ARCHIVO EXCEL MÉDICO:\n\n"
            content += f"Hojas disponibles: {', '.join(excel_file.sheet_names)}\n\n"
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                content += f"\n=== HOJA: {sheet_name} ===\n"
                content += f"Dimensiones: {df.shape[0]} registros x {df.shape[1]} variables\n"
                content += f"Variables: {', '.join(df.columns.tolist())}\n\n"
                
                # Estadísticas descriptivas
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    content += "ESTADÍSTICAS DESCRIPTIVAS:\n"
                    content += df[numeric_cols].describe().to_string() + "\n\n"
                
                # Muestra de datos
                content += "MUESTRA DE DATOS:\n"
                content += df.head(5).to_string() + "\n\n"
            
            return content
            
        except Exception as e:
            logger.error(f"❌ Error procesando Excel: {e}")
            return ""

    def intelligent_chunking(self, text: str, source: str) -> List[Dict[str, Any]]:
        """Chunking inteligente optimizado para documentos médicos"""
        # Separar por secciones médicas comunes
        medical_separators = [
            '\n\nABSTRACT\n', '\n\nINTRODUCTION\n', '\n\nMETHODS\n', 
            '\n\nRESULTS\n', '\n\nDISCUSSION\n', '\n\nCONCLUSION\n',
            '\n\nBackground\n', '\n\nObjective\n', '\n\nStudy Design\n',
            '\n\nPatients and Methods\n', '\n\nFindings\n', '\n\nInterpretation\n'
        ]
        
        chunks = []
        current_section = ""
        current_chunk = ""
        
        # Dividir por párrafos
        paragraphs = text.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # Detectar si es inicio de nueva sección médica
            is_section_header = any(sep.strip().upper() in paragraph.upper() 
                                    for sep in medical_separators)
            
            if is_section_header:
                # Guardar chunk anterior si existe
                if current_chunk:
                    chunks.append({
                        'text': current_chunk.strip(),
                        'section': current_section,
                        'tokens': len(current_chunk.split())
                    })
                current_section = paragraph
                current_chunk = paragraph + "\n\n"
            else:
                # Verificar si agregar párrafo excede el límite
                if len(current_chunk) + len(paragraph) > self.chunk_size:
                    if current_chunk:
                        chunks.append({
                            'text': current_chunk.strip(),
                            'section': current_section,
                            'tokens': len(current_chunk.split())
                        })
                    current_chunk = paragraph + "\n\n"
                else:
                    current_chunk += paragraph + "\n\n"
        
        # Agregar último chunk
        if current_chunk:
            chunks.append({
                'text': current_chunk.strip(),
                'section': current_section,
                'tokens': len(current_chunk.split())
            })
        
        # Aplicar overlapping para contexto
        final_chunks = []
        for i, chunk in enumerate(chunks):
            chunk_data = {
                'text': chunk['text'],
                'source': source,
                'chunk_index': i,
                'section': chunk['section'],
                'tokens': chunk['tokens'],
                'doc_id': f"{source}_{i}"
            }
            
            # Agregar overlap con chunk anterior
            if i > 0 and len(chunks[i-1]['text']) > 0:
                overlap_text = chunks[i-1]['text'][-self.chunk_overlap:]
                chunk_data['text'] = overlap_text + "\n\n" + chunk['text']
            
            final_chunks.append(chunk_data)
        
        return final_chunks

    def add_document(self, text: str, source: str) -> bool:
        """Agregar documento al índice FAISS"""
        try:
            # Chunking inteligente
            chunks = self.intelligent_chunking(text, source)
            
            if not chunks:
                logger.warning(f"⚠️ No se generaron chunks para: {source}")
                return False
            
            # Extraer textos para embeddings
            texts = [chunk['text'] for chunk in chunks]
            
            # Generar embeddings por lotes
            embeddings = self.get_batch_embeddings(texts)
            
            if not embeddings or len(embeddings) != len(texts):
                logger.error(f"❌ Error generando embeddings para: {source}")
                return False
            
            # Preparar embeddings para FAISS (normalizar para cosine similarity)
            embeddings_array = np.array(embeddings).astype('float32')
            faiss.normalize_L2(embeddings_array)
            
            # Agregar al índice FAISS
            start_idx = len(self.documents_metadata)
            self.index.add(embeddings_array)
            
            # Guardar metadata
            for i, chunk in enumerate(chunks):
                self.documents_metadata.append({
                    'text': chunk['text'],
                    'source': chunk['source'],
                    'chunk_index': chunk['chunk_index'],
                    'section': chunk['section'],
                    'tokens': chunk['tokens'],
                    'doc_id': chunk['doc_id'],
                    'faiss_idx': start_idx + i
                })
            
            # Guardar índice actualizado
            self._save_index()
            
            logger.info(f"✅ Documento médico procesado: {len(chunks)} chunks agregados - {source}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error agregando documento: {e}")
            return False

    def add_file(self, file_path: str, filename: str) -> bool:
        """Añadir archivo médico al sistema"""
        try:
            content = self.process_file(file_path, filename)
            
            if not content:
                logger.warning(f"⚠️ No se pudo procesar o no es contenido médico: {filename}")
                return False
            
            return self.add_document(content, filename)
            
        except Exception as e:
            logger.error(f"❌ Error añadiendo archivo {filename}: {e}")
            return False

    def search_documents(self, query: str, k: int = 10) -> List[Dict[str, Any]]:
        """Búsqueda optimizada en índice FAISS"""
        try:
            if self.index.ntotal == 0:
                logger.warning("⚠️ Índice FAISS vacío")
                return []
            
            # Generar embedding de la consulta
            query_embeddings = self.get_batch_embeddings([query])
            
            if not query_embeddings:
                logger.error("❌ Error generando embedding de consulta")
                return []
            
            query_vector = np.array([query_embeddings[0]]).astype('float32')
            faiss.normalize_L2(query_vector)
            
            # Buscar en FAISS
            scores, indices = self.index.search(query_vector, min(k, self.index.ntotal))
            
            # Preparar resultados
            results = []
            for score, idx in zip(scores[0], indices[0]):
                if idx < len(self.documents_metadata):
                    doc = self.documents_metadata[idx].copy()
                    doc['similarity_score'] = float(score)
                    results.append(doc)
            
            # Filtrar resultados por relevancia mínima
            results = [r for r in results if r['similarity_score'] > 0.5]
            
            logger.info(f"🔍 Búsqueda local: {len(results)} documentos relevantes encontrados")
            return results
            
        except Exception as e:
            logger.error(f"❌ Error en búsqueda local: {e}")
            return []

    def enhanced_medical_search(self, query: str, k: int = 10) -> List[Dict[str, Any]]:
        """Búsqueda médica mejorada: local + web médica si es necesario"""
        
        # 1. Búsqueda en documentos locales
        local_results = self.search_documents(query, k=k)
        
        # 2. Evaluar calidad de resultados locales
        high_quality_results = [doc for doc in local_results if doc['similarity_score'] > 0.7]
        
        logger.info(f"🔍 Resultados locales de alta calidad: {len(high_quality_results)}/{len(local_results)}")
        
        # 3. Si no hay suficientes resultados de calidad, buscar en web médica
        if len(high_quality_results) < 3:
            logger.info("🏥 Buscando información adicional en fuentes médicas...")
            web_results = self.medical_search(query, num_results=5)
            
            # 4. Procesar resultados web médicos
            for web_result in web_results:
                web_source = f"MED_WEB: {web_result['title']}"
                self.add_document(web_result['content'], web_source)
            
            # 5. Buscar nuevamente con información médica actualizada
            if web_results:
                local_results = self.search_documents(query, k=k)
                high_quality_results = [doc for doc in local_results if doc['similarity_score'] > 0.7]
                logger.info(f"🔍 Resultados después de búsqueda médica web: {len(high_quality_results)}")
        
        return high_quality_results or local_results[:5]

    def ask_medical_question(self, question: str, max_docs: int = 10) -> Dict[str, Any]:
        """Responder pregunta médica con evidencia científica"""
        
        # Búsqueda médica mejorada
        relevant_docs = self.enhanced_medical_search(question, k=max_docs)
        
        # Preparar contexto médico
        context = self._prepare_medical_context(relevant_docs)
        
        # Calcular nivel de confianza de la evidencia
        confidence = self._calculate_evidence_confidence(relevant_docs)
        
        # Prompt especializado para medicina
        user_prompt = f"""CONSULTA MÉDICA: {question}

        EVIDENCIA CIENTÍFICA DISPONIBLE:
        {context}

        INSTRUCCIONES ESPECÍFICAS:
        - Proporciona una respuesta basada únicamente en la evidencia científica proporcionada
        - Clasifica el nivel de evidencia cuando sea posible (Nivel I-V, Grado A-D)
        - Menciona las limitaciones de los estudios
        - Si hay información contradictoria, analízala críticamente
        - Siempre incluye la recomendación de consultar con un profesional médico
        - Estructura la respuesta para profesionales de la salud

        Responde de manera estructurada y técnicamente precisa."""
        
        try:
            response = self.openai_client.chat.completions.create(
                model=self.deployment_name,
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=2500
            )
            
            answer = response.choices[0].message.content
            
            return {
                'answer': answer,
                'sources': relevant_docs,
                'confidence': confidence,
                'query': question,
                'context_length': len(context),
                'local_sources': len([d for d in relevant_docs if not d['source'].startswith('MED_WEB:')]),
                'web_medical_sources': len([d for d in relevant_docs if d['source'].startswith('MED_WEB:')]),
                'evidence_level': self._classify_evidence_level(relevant_docs)
            }
            
        except Exception as e:
            logger.error(f"❌ Error generando respuesta médica: {e}")
            return {
                'answer': f"Error al generar respuesta médica: {str(e)}",
                'sources': relevant_docs,
                'confidence': 0.0,
                'query': question,
                'context_length': 0,
                'local_sources': 0,
                'web_medical_sources': 0,
                'evidence_level': 'Unknown'
            }

    def _prepare_medical_context(self, relevant_docs: List[Dict[str, Any]]) -> str:
        """Preparar contexto médico estructurado"""
        if not relevant_docs:
            return "No se encontraron documentos médicos relevantes."
        
        context = "=== EVIDENCIA CIENTÍFICA ===\n\n"
        
        for i, doc in enumerate(relevant_docs, 1):
            source_type = "🌐 FUENTE MÉDICA WEB" if doc['source'].startswith('MED_WEB:') else "📄 DOCUMENTO LOCAL"
            context += f"EVIDENCIA {i} ({source_type}): {doc['source']}\n"
            context += f"Sección: {doc.get('section', 'General')}\n"
            context += f"Relevancia: {doc['similarity_score']:.3f}\n"
            context += f"Tokens: {doc.get('tokens', 0)}\n"
            context += f"Contenido:\n{doc['text']}\n"
            context += "=" * 80 + "\n\n"
        
        return context

    def _calculate_evidence_confidence(self, relevant_docs: List[Dict[str, Any]]) -> float:
        """Calcular confianza basada en calidad y cantidad de evidencia médica"""
        if not relevant_docs:
            return 0.0
        
        # Factor de calidad promedio
        avg_similarity = sum(doc['similarity_score'] for doc in relevant_docs) / len(relevant_docs)
        
        # Factor de cantidad (más evidencia = más confianza, hasta un límite)
        quantity_factor = min(len(relevant_docs) / 8, 1.0)
        
        # Factor de diversidad de fuentes
        unique_sources = len(set(doc['source'] for doc in relevant_docs))
        diversity_factor = min(unique_sources / 5, 1.0)
        
        # Bonificación por fuentes médicas web confiables
        web_medical_count = len([d for d in relevant_docs if d['source'].startswith('MED_WEB:')])
        web_factor = min(web_medical_count * 0.1, 0.3)
        
        confidence = (avg_similarity * 0.4 + quantity_factor * 0.3 + 
                     diversity_factor * 0.2 + web_factor * 0.1)
        
        return min(confidence, 1.0)

    def _classify_evidence_level(self, relevant_docs: List[Dict[str, Any]]) -> str:
        """Clasificar nivel de evidencia basado en fuentes"""
        if not relevant_docs:
            return "No Evidence"
        
        source_types = []
        for doc in relevant_docs:
            source = doc['source'].lower()
            if 'pubmed' in source or 'nejm' in source or 'jama' in source or 'lancet' in source:
                source_types.append('high_impact')
            elif 'bmj' in source or 'nature' in source or 'cell' in source:
                source_types.append('peer_reviewed')
            elif 'cochrane' in source:
                source_types.append('systematic_review')
            elif 'who' in source or 'cdc' in source or 'fda' in source:
                source_types.append('regulatory')
            else:
                source_types.append('general')
        
        if 'systematic_review' in source_types:
            return "Level I (Systematic Review/Meta-analysis)"
        elif source_types.count('high_impact') >= 2:
            return "Level II (High-Quality RCT/Cohort)"
        elif 'peer_reviewed' in source_types:
            return "Level III (Peer-Reviewed Studies)"
        elif 'regulatory' in source_types:
            return "Level IV (Guidelines/Expert Opinion)"
        else:
            return "Level V (Limited Evidence)"

    def delete_medical_documents(self, sources: List[str]) -> bool:
        """Eliminar documentos médicos específicos"""
        try:
            indices_to_remove = []
            new_metadata = []
            
            # Identificar índices a remover
            for i, doc_meta in enumerate(self.documents_metadata):
                if doc_meta['source'] in sources:
                    indices_to_remove.append(doc_meta['faiss_idx'])
                else:
                    new_metadata.append(doc_meta)
            
            if not indices_to_remove:
                logger.warning("⚠️ No se encontraron documentos para eliminar")
                return False
            
            # Recrear índice FAISS sin los documentos eliminados
            if new_metadata:
                # Obtener embeddings de documentos restantes
                remaining_texts = [doc['text'] for doc in new_metadata]
                remaining_embeddings = self.get_batch_embeddings(remaining_texts)
                
                if remaining_embeddings:
                    # Crear nuevo índice
                    new_index = faiss.IndexFlatIP(self.embedding_dimension)
                    embeddings_array = np.array(remaining_embeddings).astype('float32')
                    faiss.normalize_L2(embeddings_array)
                    new_index.add(embeddings_array)
                    
                    # Actualizar índices en metadata
                    for i, doc_meta in enumerate(new_metadata):
                        doc_meta['faiss_idx'] = i
                    
                    # Reemplazar índice y metadata
                    self.index = new_index
                    self.documents_metadata = new_metadata
                else:
                    logger.error("❌ Error regenerando embeddings")
                    return False
            else:
                # Si no quedan documentos, crear índice vacío
                self.index = faiss.IndexFlatIP(self.embedding_dimension)
                self.documents_metadata = []
            
            # Guardar cambios
            self._save_index()
            
            logger.info(f"🗑️ Eliminados {len(sources)} documentos médicos")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error eliminando documentos médicos: {e}")
            return False

    def clear_medical_database(self) -> bool:
        """Limpiar completamente la base de datos médica"""
        try:
            # Crear nuevo índice vacío
            self.index = faiss.IndexFlatIP(self.embedding_dimension)
            self.documents_metadata = []
            
            # Guardar índice vacío
            self._save_index()
            
            logger.info("🗑️ Base de datos médica limpiada exitosamente")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error limpiando base de datos médica: {e}")
            return False

    def validate_medical_source(self, url_or_content: str) -> Dict[str, Any]:
        """Validar si una fuente es confiable para información médica"""
        try:
            is_url = url_or_content.startswith('http')
            
            if is_url:
                is_trusted = self._is_trusted_medical_source(url_or_content)
                domain = urlparse(url_or_content).netloc.lower()
                
                return {
                    'is_trusted': is_trusted,
                    'source_type': 'URL',
                    'domain': domain,
                    'trusted_domains': list(self.TRUSTED_MEDICAL_DOMAINS),
                    'recommendation': 'Approved' if is_trusted else 'Not Recommended'
                }
            else:
                has_medical_content = self._has_medical_content(url_or_content)
                
                return {
                    'has_medical_content': has_medical_content,
                    'source_type': 'CONTENT',
                    'medical_keywords_found': [kw for kw in self.MEDICAL_KEYWORDS if kw in url_or_content.lower()],
                    'recommendation': 'Approved' if has_medical_content else 'Review Required'
                }
                
        except Exception as e:
            logger.error(f"❌ Error validando fuente médica: {e}")
            return {
                'error': str(e),
                'recommendation': 'Error in Validation'
            }